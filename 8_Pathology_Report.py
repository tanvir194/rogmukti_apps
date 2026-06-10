import streamlit as st
import sqlite3
import json
import os
from docx import Document
from datetime import datetime

# সিকিউরিটি চেক
if "logged_in" not in st.session_state or not st.session_state.logged_in:
    st.warning("🔒 অ্যাক্সেস রিফিউজড! দয়া করে আগে মেইন পেজ থেকে লগইন করুন।")
    st.stop()

st.title("🔬 প্যাথলজি রিপোর্ট (কাস্টম ওয়ার্ড ফরম্যাট)")
st.write("---")

# ডাটাবেজ কানেকশন
conn = sqlite3.connect("rogmukti_clinic_fix.db")
c = conn.cursor()
c.execute("""
CREATE TABLE IF NOT EXISTS pathology_reports (
    invoice_id INTEGER PRIMARY KEY,
    report_data TEXT,
    reported_date TEXT
)
""")
conn.commit()
conn.close()

# --- ১. বিল নম্বর দিয়ে রোগী খোঁজা ও ডাটা এন্ট্রি ---
search_id = st.number_input("বিল নম্বর (Invoice ID) দিয়ে রোগীর টেস্ট খুঁজুন:", min_value=0, step=1, value=0)

if search_id > 0:
    conn = sqlite3.connect("rogmukti_clinic_fix.db")
    c = conn.cursor()
    c.execute("SELECT patient_name, age, phone, doctor, selected_tests, billing_date FROM billing_records WHERE id = ?", (search_id,))
    patient_row = c.fetchone()
    
    c.execute("SELECT report_data FROM pathology_reports WHERE invoice_id = ?", (search_id,))
    saved_report_row = c.fetchone()
    conn.close()
    
    if patient_row:
        p_name, p_age, p_phone, p_doctor, p_tests_str, p_date = patient_row
        st.success(f"🔍 রোগী পাওয়া গেছে: {p_name} | ডাক্তার: {p_doctor}")
        
        raw_tests = [item.strip() for item in p_tests_str.split('|') if item.strip()]
        tests_to_process = []
        for item in raw_tests:
            if ":" in item:
                tests_to_process.append(item.split(":", 1))
            else:
                tests_to_process.append(item)
                
        saved_data = {}
        if saved_report_row:
            try: saved_data = json.loads(saved_report_row)
            except: pass

        st.subheader("📝 টেস্টের ফলাফল লিখুন")
        new_report_data = {}
        
        for test in tests_to_process:
            st.markdown(f"##### 🧪 টেস্ট: **{test}**")
            col1, col2, col3 = st.columns(3)
            with col1:
                result = st.text_input(f"ফলাফল (Result) - {test}:", value=saved_data.get(test, {}).get("result", ""), key=f"res_{test}")
            with col2:
                unit = st.text_input(f"ইউনিট (Unit) - {test}:", value=saved_data.get(test, {}).get("unit", ""), key=f"unit_{test}")
            with col3:
                ref_range = st.text_input(f"নরমাল রেঞ্জ (Reference) - {test}:", value=saved_data.get(test, {}).get("ref", ""), key=f"ref_{test}")
                
            new_report_data[test] = {"result": result, "unit": unit, "ref": ref_range}
            st.markdown("---")
            
        if st.button("💾 রিপোর্টের ডাটা সেভ করুন", type="primary", use_container_width=True):
            conn = sqlite3.connect("rogmukti_clinic_fix.db")
            c = conn.cursor()
            json_str = json.dumps(new_report_data)
            current_time_str = datetime.now().strftime("%Y-%m-%d")
            c.execute("""
                INSERT INTO pathology_reports (invoice_id, report_data, reported_date)
                VALUES (?, ?, ?)
                ON CONFLICT(invoice_id) DO UPDATE SET report_data = ?, reported_date = ?
            """, (search_id, json_str, current_time_str, json_str, current_time_str))
            conn.commit()
            conn.close()
            st.success("✅ ডাটা সেভ হয়েছে! এবার নিচে আপনার প্যাড আপলোডের অপশন চালু হয়েছে।")
            st.rerun()

        # --- ২. কাস্টম ওয়ার্ড ফরম্যাট আপলোড ও মেগা ডাউনলোড লজিক ---
        if saved_report_row:
            st.write("---")
            st.subheader("📁 এই রিপোর্টের জন্য আপনার MS Word ফরম্যাট ফাইলটি দিন")
            
            # 🌟 জাদুকরী পরিবর্তন: এখানে doc এবং docx দুইটাই টাইপ হিসেবে দিয়ে দেওয়া হয়েছে
            uploaded_template = st.file_uploader(
                "আপনার মোবাইল থেকে টেস্টের ফাইলটি সিলেক্ট করুন:", 
                type=["doc", "docx"], 
                key="doc_uploader"
            )
            
            if uploaded_template is not None:
                try:
                    # ফাইল নেম এক্সটেনশন চেক
                    file_extension = os.path.splitext(uploaded_template.name).lower()
                    
                    # বাইনারি ডাটা থেকে ডকুমেন্ট ওপেন করা
                    doc = Document(uploaded_template)
                    
                    # প্যারাগ্রাফের ট্যাগ পরিবর্তন
                    for paragraph in doc.paragraphs:
                        if "{{PATIENT_NAME}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{PATIENT_NAME}}", str(p_name))
                        if "{{INVOICE_ID}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{INVOICE_ID}}", f"#{search_id}")
                        if "{{AGE}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{AGE}}", str(p_age))
                        if "{{DOCTOR}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{DOCTOR}}", str(p_doctor))
                        if "{{DATE}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{DATE}}", str(p_date))
                    
                    # টেবিলের ভেতরের ট্যাগ পরিবর্তন
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if "{{PATIENT_NAME}}" in cell.text: cell.text = cell.text.replace("{{PATIENT_NAME}}", str(p_name))
                                if "{{INVOICE_ID}}" in cell.text: cell.text = cell.text.replace("{{INVOICE_ID}}", f"#{search_id}")
                                if "{{AGE}}" in cell.text: cell.text = cell.text.replace("{{AGE}}", str(p_age))
                                if "{{DOCTOR}}" in cell.text: cell.text = cell.text.replace("{{DOCTOR}}", str(p_doctor))
                                if "{{DATE}}" in cell.text: cell.text = cell.text.replace("{{DATE}}", str(p_date))
                                
                                if "{{TEST_RESULTS}}" in cell.text:
                                    cell.text = "" 
                                    results_text = ""
                                    for t_name, t_val in saved_data.items():
                                        results_text += f"{t_name} \t Result: {t_val['result']} \t Unit: {t_val['unit']} \t Ref: {t_val['ref']}\n"
                                    cell.text = results_text

                    # রোগীর আইডি অনুযায়ী আউটপুট ফাইল তৈরি
                    output_filename = f"Final_Report_ID_{search_id}{file_extension}"
                    doc.save(output_filename)
                    
                    with open(output_filename, "rb") as file:
                        st.download_button(
                            label="📥 তৈরি হওয়া ফাইনাল রিপোর্ট ফাইলটি ডাউনলোড করুন",
                            data=file,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if file_extension == ".docx" else "application/msword",
                            use_container_width=True
                        )
                    
                    if os.path.exists(output_filename):
                        os.remove(output_filename)
                        
                except Exception as e:
                    st.error(f"❌ এই ফাইলটি সরাসরি প্রসেস করতে সমস্যা হচ্ছে। মাইক্রোসফট ওয়ার্ডের অনেক পুরনো সংস্করণের ফাইল হলে এই সমস্যা হতে পারে।")
                    st.info("💡 বিকল্প করণীয়: যদি এরর দেখায়, তবে ফাইলটিকে WPS Office দিয়ে একবার ওপেন করে জাস্ট 'Save As' করে নতুন নামে সেভ করে নিয়ে ট্রাই করুন।")
            else:
                st.info("💡 ডাটা সেভ করা আছে। ওপরে আপনার মোবাইল থেকে এই টেস্টের নির্দিষ্ট ফাইলটি (যেমন: Widal.doc) সিলেক্ট করলে ডাউনলোড বাটন চলে আসবে।")
    else:
        st.error("⚠️ এই বিল নম্বরের কোনো রোগীর তথ্য পাওয়া যায়নি।")
