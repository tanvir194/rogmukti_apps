import streamlit as st
import sqlite3
import json
import os
from docx import Document
from datetime import datetime

# সিকিউরিটি চেক
if "logged_in" not in st.session_state or not st.session_state.logged_in:
    st.warning("🔒 অ্যাক্সেস রিফিউজড! দয়া করে আগে মেইন পেজ থেকে লগইন করুন।")
    st.stop()

st.title("🔬 প্যাথলজি রিপোর্ট (কাস্টম ওয়ার্ড ফরম্যাট)")
st.write("---")

# 🌟 ক্লাউড ডাটাবেজ এরর সমাধানের জন্য নিখুঁত ডিরেক্টরি পাথ লজিক
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DB_PATH = os.path.join(BASE_DIR, "rogmukti_clinic_fix.db")
TEMPLATE_DIR = os.path.join(BASE_DIR, "report_templates")
TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "pathology_template.docx")

if not os.path.exists(TEMPLATE_DIR):
    os.makedirs(TEMPLATE_DIR)

# ডাটাবেজ কানেকশন ও টেবিল তৈরি নিশ্চিত করা
try:
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
    CREATE TABLE IF NOT EXISTS pathology_reports (
        invoice_id INTEGER PRIMARY KEY,
        report_data TEXT,
        reported_date TEXT
    )
    """)
    conn.commit()
    conn.close()
except Exception as e:
    st.error(f"❌ ডাটাবেজ কানেকশন তৈরি করা যাচ্ছে না: {e}")

# --- ১. মাইক্রোসফট ওয়ার্ড ফরম্যাট আপলোড সেকশন ---
st.subheader("📁 আপনার অফিশিয়াল MS Word (.doc/.docx) ফরম্যাট আপলোড করুন")
uploaded_template = st.file_uploader("ল্যাবের প্যাড বা ফরম্যাট ফাইলটি এখানে একবার আপলোড করে রাখুন:", type=["doc", "docx"])

if uploaded_template is not None:
    with open(TEMPLATE_PATH, "wb") as f:
        f.write(uploaded_template.getbuffer())
    st.success("✅ আপনার কাস্টম ওয়ার্ড ফরম্যাটটি সফলভাবে আপলোড ও সেভ হয়েছে!")

# ফরম্যাট ফাইলটি আপলোড করা আছে কি না চেক
if not os.path.exists(TEMPLATE_PATH):
    st.info("ℹ️ শুরু করার জন্য ওপরে আপনার ল্যাবের রিসিট বা রিপোর্টের একটি মাইক্রোসফট ওয়ার্ড ফরম্যাট আপলোড করুন।")
    st.markdown("""
    💡 **ওয়ার্ড ফাইলে লেখার নিয়ম:** আপনার ওয়ার্ড ফাইলের যেখানে রোগীর নাম, বয়স বসাতে চান, সেখানে নিচের ট্যাগগুলো হুবহু লিখে রাখুন:
    * রোগীর নামের জায়গায় লিখুন: `{{PATIENT_NAME}}`
    * ইনভয়েস আইডির জায়গায় লিখুন: `{{INVOICE_ID}}`
    * বয়স: `{{AGE}}` | ডাক্তার: `{{DOCTOR}}` | তারিখ: `{{DATE}}`
    * টেস্টের ফলাফল টেবিলের ভেতরে যেখানে টেস্টের নাম ও রেজাল্ট বসবে সেখানে লিখুন: `{{TEST_RESULTS}}`
    """)
    st.stop()
else:
    st.caption("✅ ল্যাবের নিজস্ব ওয়ার্ড ফরম্যাটটি ব্যাকগ্রাউন্ডে রেডি আছে।")

st.write("---")

# --- ২. বিল নম্বর দিয়ে রোগী খোঁজা ও ডাটা এন্ট্রি ---
search_id = st.number_input("বিল নম্বর (Invoice ID) দিয়ে রোগীর টেস্ট খুঁজুন:", min_value=0, step=1, value=0)

if search_id > 0:
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT patient_name, age, phone, doctor, selected_tests, billing_date FROM billing_records WHERE id = ?", (search_id,))
        patient_row = c.fetchone()
        
        c.execute("SELECT report_data FROM pathology_reports WHERE invoice_id = ?", (search_id,))
        saved_report_row = c.fetchone()
        conn.close()
        
        if patient_row:
            p_name, p_age, p_phone, p_doctor, p_tests_str, p_date = patient_row
            st.success(f"🔍 রোগী পাওয়া গেছে: {p_name} | ডাক্তার: {p_doctor}")
            
            # টেস্ট আলাদা করা
            raw_tests = [item.strip() for item in p_tests_str.split('|') if item.strip()]
            tests_to_process = []
            for item in raw_tests:
                if ":" in item:
                    tests_to_process.append(item.split(":", 1))
                else:
                    tests_to_process.append(item)
                    
            saved_data = {}
            if saved_report_row:
                try:
                    saved_data = json.loads(saved_report_row)
                except:
                    pass

            st.subheader("📝 টেস্টের ফলাফল ও রেফারেন্স ভ্যালু লিখুন")
            new_report_data = {}
            
            for test in tests_to_process:
                st.markdown(f"##### 🧪 টেস্ট: **{test}**")
                col1, col2, col3 = st.columns(3)
                with col1:
                    result = st.text_input(f"ফলাফল (Result) - {test}:", value=saved_data.get(test, {}).get("result", ""), key=f"res_{test}")
                with col2:
                    unit = st.text_input(f"ইউনিট (Unit) - {test}:", value=saved_data.get(test, {}).get("unit", ""), key=f"unit_{test}")
                with col3:
                    ref_range = st.text_input(f"নরমাল রেঞ্জ (Reference) - {test}:", value=saved_data.get(test, {}).get("ref", ""), key=f"ref_{test}")
                    
                new_report_data[test] = {"result": result, "unit": unit, "ref": ref_range}
                st.markdown("---")
                
            if st.button("💾 রিপোর্টের ডাটা সেভ করুন", type="primary", use_container_width=True):
                conn = sqlite3.connect(DB_PATH)
                c = conn.cursor()
                json_str = json.dumps(new_report_data)
                current_time_str = datetime.now().strftime("%Y-%m-%d")
                
                c.execute("""
                    INSERT INTO pathology_reports (invoice_id, report_data, reported_date)
                    VALUES (?, ?, ?)
                    ON CONFLICT(invoice_id) DO UPDATE SET report_data = ?, reported_date = ?
                """, (search_id, json_str, current_time_str, json_str, current_time_str))
                conn.commit()
                conn.close()
                st.success("✅ ডাটা সেভ হয়েছে! নিচে ফাইল ডাউনলোড অপশন চালু হয়েছে।")
                st.rerun()

            # --- ৩. ওয়ার্ড ফাইল জেনারেট ও ডাউনলোড লজিক ---
            if saved_report_row:
                st.subheader("📥 আপনার কাস্টমাইজড ওয়ার্ড রিপোর্টটি ডাউনলোড করুন")
                
                try:
                    doc = Document(TEMPLATE_PATH)
                    
                    # প্যারাগ্রাফের ট্যাগ পরিবর্তন করা
                    for paragraph in doc.paragraphs:
                        if "{{PATIENT_NAME}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{PATIENT_NAME}}", str(p_name))
                        if "{{INVOICE_ID}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{INVOICE_ID}}", f"#{search_id}")
                        if "{{AGE}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{AGE}}", str(p_age))
                        if "{{DOCTOR}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{DOCTOR}}", str(p_doctor))
                        if "{{DATE}}" in paragraph.text: paragraph.text = paragraph.text.replace("{{DATE}}", str(p_date))
                    
                    # টেবিলের ভেতরের ট্যাগ পরিবর্তন
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if "{{PATIENT_NAME}}" in cell.text: cell.text = cell.text.replace("{{PATIENT_NAME}}", str(p_name))
                                if "{{INVOICE_ID}}" in cell.text: cell.text = cell.text.replace("{{INVOICE_ID}}", f"#{search_id}")
                                if "{{AGE}}" in cell.text: cell.text = cell.text.replace("{{AGE}}", str(p_age))
                                if "{{DOCTOR}}" in cell.text: cell.text = cell.text.replace("{{DOCTOR}}", str(p_doctor))
                                if "{{DATE}}" in cell.text: cell.text = cell.text.replace("{{DATE}}", str(p_date))
                                
                                if "{{TEST_RESULTS}}" in cell.text:
                                    cell.text = "" 
                                    results_text = ""
                                    for t_name, t_val in saved_data.items():
                                        results_text += f"{t_name} \t\t Result: {t_val['result']} \t Unit: {t_val['unit']} \t Ref: {t_val['ref']}\n"
                                    cell.text = results_text

                    output_path = os.path.join(BASE_DIR, f"Report_{search_id}.docx")
                    doc.save(output_path)
                    
                    with open(output_path, "rb") as file:
                        st.download_button(
                            label="📥 জেনারেটেড MS Word (.docx) ফাইল ডাউনলোড করুন",
                            data=file,
                            file_name=f"Pathology_Report_ID_{search_id}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    
                    if os.path.exists(output_path):
                        os.remove(output_path)
                        
                except Exception as e:
                    st.error(f"❌ ওয়ার্ড ফাইলটি প্রসেস করতে সমস্যা হচ্ছে। এরর: {e}")
        else:
            st.error(f"⚠️ এই বিল নম্বরের ({search_id}) কোনো রোগীর তথ্য পাওয়া যায়নি। অনুগ্রহ করে '1_Patient_Entry.py' পেজে গিয়ে নিশ্চিত হোন যে এই নম্বরের বিলটি ডাটাবেজে সঠিক নামে সাবমিট করা আছে কি না।")
    except Exception as e:
        st.error(f"❌ ডাটাবেজ অপারেশন ব্যর্থ হয়েছে। এরর টেক্সট: {e}")
